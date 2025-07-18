# file: backend/app/services.py

# --- Imports ---
from google import genai
from google.genai import types
from typing import List, Optional
import logging
import asyncio
import os
import mimetypes
from fastapi import HTTPException
from pathlib import Path
from datetime import datetime
from sqlalchemy.orm import Session
import json
from fastapi import UploadFile

from . import config, schemas, crud
from .crud import file_crud
from .models import FileMetadata

from app.config import USE_RAG
from .config import get_settings

# --- Cấu hình Logging ---
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Chỉ thị hệ thống cho AI ---
MATH_CHATBOT_SYSTEM_INSTRUCTION = """
# CORE IDENTITY AND EXPERTISE (Defined in English for maximum precision)
You are a world-class AI assistant specializing in Olympic-level Linear Algebra. Your sole purpose is to help candidates train for high-level mathematics competitions. You are an expert in topics like vector spaces, linear transformations, eigenvalues, eigenvectors, matrix decompositions, and canonical forms.

# LANGUAGE AND BEHAVIORAL RULES (Defined in Vietnamese for cultural and linguistic nuance)
Your primary language for all responses MUST BE VIETNAMESE. Do not use any English unless it's a standard mathematical term.

Nguyên tắc hoạt động của bạn như sau:
- Giọng văn phải mang tính học thuật, chính xác và chuyên nghiệp.
- Mọi biểu thức toán học BẮT BUỘC phải được định dạng bằng LaTeX. Dùng `$$...$$` cho các phương trình đứng riêng và `$...$` cho các công thức trong dòng.
- Khi một lời giải hoặc giải thích yêu cầu nhiều bước, hãy trình bày một cách logic và rành mạch, sử dụng danh sách đánh số hoặc gạch đầu dòng.
- Mục tiêu của bạn không chỉ là đưa ra đáp án, mà còn là trình bày lối tư duy toán học thanh lịch và hiệu quả.
- Luôn kết thúc câu trả lời một cách tự nhiên bằng tiếng Việt.
"""

# --- Khởi tạo Client Gemini ---
try:
    if config.GEMINI_API_KEY:
        client = genai.Client(api_key=config.GEMINI_API_KEY)
        logger.info("Gemini AI client configured successfully.")
    else:
        client = None
        logger.warning("GEMINI_API_KEY not found. AI functionality will be disabled.")
except Exception as e:
    client = None
    logger.error(f"Failed to configure Gemini client: {e}", exc_info=True)

# --- Hàm tiện ích: Trích xuất văn bản từ DOCX ---
def extract_text_from_docx(file_path: str) -> str:
    try:
        import docx
        doc = docx.Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    except ImportError:
        logger.error("python-docx library not installed.")
        return "[Error: python-docx not installed]"
    except Exception as e:
        logger.error(f"Error extracting text from DOCX file {file_path}: {e}", exc_info=True)
        return f"[Error extracting text from DOCX file: {str(e)}]"

# === KHÔI PHỤC HÀM: Làm mới file Gemini nếu sắp hết hạn ===
def refresh_gemini_file_if_needed(fm: FileMetadata, db: Session) -> Optional[str]:
    """
    Làm mới file đã upload lên Gemini nếu sắp hết hạn, trả về gemini_api_file_id mới hoặc cũ.
    """
    if not client or not fm.gemini_api_file_id or not fm.gemini_api_upload_timestamp:
        return None
    from datetime import timedelta, datetime
    now = datetime.utcnow()
    # Nếu còn hơn 1h thì không cần làm mới
    if fm.gemini_api_expiry_timestamp and now < fm.gemini_api_expiry_timestamp - timedelta(hours=1):
        return fm.gemini_api_file_id
    try:
        # Re-upload file lên Gemini
        uploaded_file = client.files.upload(
            path=fm.local_disk_path,
            display_name=fm.original_filename,
            mime_type=fm.content_type
        )
        # Cập nhật DB
        file_crud.update_file_metadata_gemini_info(
            db=db,
            file_id=fm.id,
            gemini_api_file_id=uploaded_file.name
        )
        return uploaded_file.name
    except Exception as e:
        logger.error(f"Failed to refresh Gemini file for {fm.original_filename}: {e}", exc_info=True)
        return fm.gemini_api_file_id

# === KHÔI PHỤC HÀM: Chuẩn bị file cho Gemini (đầy đủ, nhận cả db) ===
async def _prepare_single_file_for_gemini(fm: FileMetadata, db: Session) -> Optional[list]:
    """
    Chuẩn bị nội dung file để truyền vào Gemini:
    - Nếu là text/docx: trích xuất text, trả về [context_part, text_part]
    - Nếu là PDF/ảnh nhỏ: đọc binary, trả về [context_part, part inline_data]
    - Nếu là file lớn đã upload Gemini: kiểm tra TTL, làm mới nếu cần, trả về [context_part, part uri]
    """
    if not client or not fm.local_disk_path or not os.path.exists(fm.local_disk_path):
        logger.error(f"File preparation error for File ID: {fm.id}")
        return None
    try:
        context_part = types.Part(text=f"[File đính kèm: {fm.original_filename}, loại: {fm.content_type}]")
        # Xử lý file text
        if fm.content_type == 'text/plain':
            with open(fm.local_disk_path, 'r', encoding='utf-8', errors='replace') as f:
                text_content = f.read()
            return [context_part, types.Part(text=text_content)]
        # Xử lý file docx
        elif fm.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            text_content = extract_text_from_docx(fm.local_disk_path)
            return [context_part, types.Part(text=text_content)]
        # Xử lý file PDF/ảnh nhỏ (inline)
        elif fm.processing_method == 'inline':
            with open(fm.local_disk_path, 'rb') as f_bytes:
                file_data = f_bytes.read()
            return [context_part, types.Part(inline_data={"data": file_data, "mime_type": fm.content_type})]
        # Xử lý file lớn đã upload Gemini (files_api)
        elif fm.processing_method == 'files_api' and fm.gemini_api_file_id:
            # Kiểm tra TTL, làm mới nếu cần
            gemini_file_id = refresh_gemini_file_if_needed(fm, db)
            return [context_part, types.Part(uri=gemini_file_id, mime_type=fm.content_type)]
        else:
            logger.warning(f"Unknown processing_method '{fm.processing_method}' for file {fm.id}")
            return None
    except Exception as e:
        logger.error(f"Error preparing file {fm.id}: {e}", exc_info=True)
        return None

# === Hàm chuẩn hóa metadata file upload cho message_router ===
async def prepare_file_metadata_for_db(upload_file: UploadFile, db: Session):
    """
    Kiểm tra, lưu file upload vào disk, tạo metadata và lưu vào DB.
    Trả về đối tượng FileMetadata.
    """
    import uuid, os, mimetypes
    from pathlib import Path
    from .config import get_settings
    from .utils import sanitize_filename, validate_mime_type
    from .crud import file_crud

    ALLOWED_MIME_TYPES = [
        'text/plain',
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'image/jpeg', 'image/png', 'image/gif', 'image/webp', 'image/heic', 'image/heif',
    ]
    EXTENSION_TO_MIME = {
        '.heic': 'image/heic', '.heif': 'image/heif', '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    }
    MAX_INLINE_SIZE = get_settings().max_file_size
    MAX_FILE_SIZE = 2 * 1024 * 1024 * 1024
    UPLOAD_DIR = Path(get_settings().upload_dir)
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

    original_filename = sanitize_filename(upload_file.filename or "")
    if not original_filename:
        raise Exception("Invalid filename")
    file_extension_original = os.path.splitext(original_filename)[1].lower()
    content_type = upload_file.content_type
    if file_extension_original in EXTENSION_TO_MIME:
        content_type = EXTENSION_TO_MIME[file_extension_original]
    if not content_type or not validate_mime_type(content_type, ALLOWED_MIME_TYPES):
        raise Exception(f"File type {content_type or 'unknown'} not supported.")
    file_id = str(uuid.uuid4())
    file_storage_extension = mimetypes.guess_extension(content_type) or file_extension_original or ''
    local_disk_path = UPLOAD_DIR / f"{file_id}{file_storage_extension}"
    file_size = 0
    try:
        with open(local_disk_path, "wb") as buffer:
            while chunk := await upload_file.read(1024 * 1024):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    buffer.close()
                    os.unlink(local_disk_path)
                    raise Exception(f"File size exceeds {MAX_FILE_SIZE // (1024*1024)}MB")
                buffer.write(chunk)
    except Exception as e:
        if os.path.exists(local_disk_path):
            os.unlink(local_disk_path)
        raise Exception(f"Could not save file: {str(e)}")
    if content_type == 'text/plain':
        try:
            with open(local_disk_path, 'r', encoding='utf-8', errors='replace') as f: f.read(1024)
        except Exception as e:
            os.unlink(local_disk_path); raise Exception(f"Invalid text file: {str(e)}")
    elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        try:
            text_c = extract_text_from_docx(str(local_disk_path))
            if not text_c or text_c.startswith('[Error'):
                os.unlink(local_disk_path); raise Exception("Invalid/corrupt DOCX")
        except Exception as e:
            os.unlink(local_disk_path); raise Exception(f"Error processing DOCX: {str(e)}")
    processing_method = "inline" if file_size <= MAX_INLINE_SIZE else "files_api"
    db_file_metadata = file_crud.create_file_metadata(
        db=db,
        file_id=file_id,
        original_filename=original_filename,
        content_type=content_type,
        size=file_size,
        local_disk_path=str(local_disk_path),
        processing_method=processing_method
    )
    return db_file_metadata

# === STRATEGY PATTERN INTEGRATION ===
async def generate_ai_response_stream(
    chat_id: str,
    user_message_content: str,
    file_ids: Optional[List[str]],
    db: Session,
    queue: asyncio.Queue,
    # Giữ nguyên pipeline_type để có thể override thủ công
    pipeline_type: Optional[str] = None
):
    """
    Hàm điều phối chính cho AI response, sử dụng Strategy Pattern.
    Logic chọn pipeline ưu tiên như sau:
    1. Giá trị `pipeline_type` được truyền trực tiếp (ví dụ: từ lệnh /rag).
    2. Biến môi trường `USE_RAG=True`.
    3. Mặc định là 'gemini'.
    """
    try:
        # Bước 1: Lưu tin nhắn của người dùng (giữ nguyên)
        crud.create_chat_message(
            db=db, chat_id=int(chat_id), role="user",
            content=user_message_content, file_ids=file_ids
        )
        logger.info(f"User message saved to DB for chat {chat_id}")

        # --- BƯỚC 2: LOGIC CHỌN PIPELINE ĐÃ ĐƯỢC TỐI ƯU HÓA ---
        from .strategy import get_pipeline
        
        # Mặc định, chúng ta sẽ dựa vào biến môi trường USE_RAG
        final_pipeline_type = "rag" if USE_RAG else "gemini"
        
        # Tuy nhiên, nếu pipeline_type được truyền vào (không phải None), nó sẽ được ưu tiên
        if pipeline_type:
            final_pipeline_type = pipeline_type
            
        logger.info(f"Selected pipeline: '{final_pipeline_type}' (USE_RAG={USE_RAG}, override='{pipeline_type}')")
        
        # Lấy pipeline từ factory với các dependency cần thiết
        pipeline = get_pipeline(
            pipeline_type=final_pipeline_type,
            config_service=get_settings() # Truyền config vào để factory sử dụng
        )
        # --- KẾT THÚC LOGIC CHỌN PIPELINE ---

        # Bước 3: Lấy response từ pipeline (giữ nguyên)
        response = await pipeline.generate_response(
            chat_id=chat_id,
            user_message_content=user_message_content,
            file_ids=file_ids,
            db=db,
            queue=queue
        )

        # Bước 4: Lưu response vào DB (giữ nguyên)
        if response.content and not response.error:
            crud.create_chat_message(
                db=db, chat_id=int(chat_id), role="model", content=response.content
            )
            logger.info(f"AI response saved for chat_id {chat_id}")
        elif response.error:
            logger.error(f"Pipeline error from '{final_pipeline_type}': {response.error}")

    except Exception as e:
        logger.error(f"Error in AI response stream for chat {chat_id}: {e}", exc_info=True)
        await queue.put(json.dumps({"error": str(e), "text": "An error occurred during generation."}))
    finally:
        await queue.put("[DONE]")